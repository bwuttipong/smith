import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_element(name):
    return OxmlElement(name)

def add_p_border_bottom(p, color_hex="1F4E79", size=12):
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = create_element('w:pBdr')
        pPr.append(pBdr)
    bottom = create_element('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)

def build_resume():
    doc = Document()
    
    # Configure margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Colors
    color_primary = RGBColor(31, 78, 121)    # Dark Blue Accent (#1F4E79)
    color_text_dark = RGBColor(51, 51, 51)   # Off-black charcoal
    color_text_muted = RGBColor(102, 102, 102) # Grey

    # Define base formatting
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = color_text_dark

    # Document Header - Name
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_after = Pt(2)
    run_name = p_name.add_run("NATTAWAT THONMON")
    run_name.font.size = Pt(22)
    run_name.font.bold = True
    run_name.font.color.rgb = color_primary

    # Subheader - Contact Information
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(4)
    run_contact = p_contact.add_run(
        "Address: 99/26 Soi 47, Bang-sai, Muang Chonburi, Chonburi 200000\n"
        "Telephone: 088-1952096  |  Email: khayombath@hotmail.com  |  Birth Date: 09 Nov 1986"
    )
    run_contact.font.size = Pt(9.5)
    run_contact.font.color.rgb = color_text_muted

    # Subheader - Additional details (Expected Salary & Start Date)
    p_details = doc.add_paragraph()
    p_details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_details.paragraph_format.space_after = Pt(18)
    run_details = p_details.add_run(
        "Expected Salary: 33,000 THB / Month (or Negotiable)  |  Availability: 30-Day Notice"
    )
    run_details.font.size = Pt(9.5)
    run_details.font.italic = True
    run_details.font.color.rgb = color_text_muted

    # Helper function to add a section header
    def add_section_header(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(title.upper())
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = color_primary
        add_p_border_bottom(p, color_hex="1F4E79", size=8)
        return p

    # 1. Career Objective
    add_section_header("Career Objective")
    p_obj = doc.add_paragraph()
    p_obj.paragraph_format.space_after = Pt(10)
    p_obj.paragraph_format.line_spacing = 1.15
    run_obj = p_obj.add_run(
        "Seeking a challenging position in Warehouse & Logistics management to utilize my extensive "
        "experience in inventory control, receiving/put-away operations, delivery truck coordination, "
        "and process optimization (Kaizen / PDCA) to drive efficiency and operational excellence."
    )

    # 2. Work Experience
    add_section_header("Work Experience")

    # Helper function for experience headers (Title & Date using a single-row borderless table)
    def add_job_header(title, date_str, company=""):
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False
        
        # Set widths
        table.columns[0].width = Inches(5.0)
        table.columns[1].width = Inches(1.5)
        
        cell_left = table.cell(0, 0)
        cell_right = table.cell(0, 1)
        
        # Format left cell (Job Title and Company)
        p_left = cell_left.paragraphs[0]
        p_left.paragraph_format.space_before = Pt(6)
        p_left.paragraph_format.space_after = Pt(2)
        
        run_title = p_left.add_run(title)
        run_title.font.bold = True
        run_title.font.size = Pt(11)
        run_title.font.color.rgb = color_primary
        
        if company:
            run_comp = p_left.add_run(f"  |  {company}")
            run_comp.font.italic = True
            run_comp.font.size = Pt(10)
            run_comp.font.color.rgb = color_text_dark
            
        # Format right cell (Date)
        p_right = cell_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_right.paragraph_format.space_before = Pt(6)
        p_right.paragraph_format.space_after = Pt(2)
        
        run_date = p_right.add_run(date_str)
        run_date.font.bold = True
        run_date.font.size = Pt(10)
        run_date.font.color.rgb = color_text_muted

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = color_text_dark
        return p

    # --- Job 1 ---
    add_job_header("Warehouse & Production Control (WH/PC) Leader", "June 2023 - Present", "San-en (Thailand) Co., Ltd.")
    add_bullet("Responsible for warehouse store operations and the picking team.")
    add_bullet("Control and monitor receiving, production, and delivery plans, including loading operations and incoming shipments.")
    add_bullet("Report manual packaging and stock levels directly to the manager.")
    add_bullet("Lead safety briefings and maintain safety communications across teams to ensure compliance.")
    add_bullet("Actively participate in Kaizen and continuous process improvement teams.")
    add_bullet("Completed the ISO 9001:2015 Internal Quality Audit training program (October 2025).")
    add_bullet("Awarded Employee of the Year (2024) for outstanding performance.")

    # --- Job 2 ---
    add_job_header("Warehouse Foreman (WH Foreman)", "2020 - June 2023", "Oizuru Chugen Packaging System Co., Ltd.")
    
    p_yr_2023 = doc.add_paragraph()
    p_yr_2023.paragraph_format.space_before = Pt(4)
    p_yr_2023.paragraph_format.space_after = Pt(2)
    p_yr_2023.paragraph_format.left_indent = Inches(0.25)
    run_yr_2023 = p_yr_2023.add_run("Year 2023 Duties:")
    run_yr_2023.font.bold = True
    run_yr_2023.font.size = Pt(9.5)
    run_yr_2023.font.color.rgb = color_primary

    add_bullet("Led and managed the Receiving, Put-away, and Inventory control teams.")
    add_bullet("Prepared and presented key operational data for weekly manager meetings.")
    add_bullet("Generated, analyzed, and maintained inventory reports to ensure high accuracy.")
    add_bullet("Supported ISO documentation compliance and audit requirements.")
    add_bullet("Tracked, monitored, and achieved departmental Key Performance Indicators (KPIs).")

    p_yr_2022 = doc.add_paragraph()
    p_yr_2022.paragraph_format.space_before = Pt(4)
    p_yr_2022.paragraph_format.space_after = Pt(2)
    p_yr_2022.paragraph_format.left_indent = Inches(0.25)
    run_yr_2022 = p_yr_2022.add_run("Year 2022 Duties:")
    run_yr_2022.font.bold = True
    run_yr_2022.font.size = Pt(9.5)
    run_yr_2022.font.color.rgb = color_primary

    add_bullet("Managed the Receiving, Put-away, and Inventory control teams, fostering a multi-skilled workforce.")
    add_bullet("Supported warehouse project improvements focusing on space utilization (sourcing suppliers for rack installations and coordinating projects through to successful completion).")
    add_bullet("Updated and modernized departmental job descriptions (JDs) for 2023 in collaboration with the Assistant Manager.")
    add_bullet("Managed departmental reports and supported all logistics and warehouse KPIs.")

    # --- Job 2 Cont'd (2021) ---
    add_job_header("Warehouse Foreman (WH Foreman) - Continued", "2021", "Oizuru Chugen Packaging System Co., Ltd.")
    add_bullet("Responsible for the Delivery Truck Control Plan, coordinating routing and truck dispatching daily.")
    add_bullet("Monitored daily targets and performance against the plan, providing daily summary reports to the Manager.")
    add_bullet("Evaluated transportation metrics monthly and updated Work Instructions (WI) for the delivery team to improve efficiency.")
    add_bullet("Conducted monthly visual audit checks on transport trucks to ensure safety and compliance.")
    add_bullet("Assisted the Supervisor/Manager in KPI management, successfully keeping transportation costs under the target of 1.50% of total sales (actual achieved: 1.77% due to unplanned orders).")
    
    def add_pdca_detail(label, items, is_issue=True):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run_lbl = p.add_run(f"{label}:\n")
        run_lbl.font.bold = True
        run_lbl.font.size = Pt(9.5)
        run_lbl.font.color.rgb = RGBColor(180, 50, 50) if is_issue else RGBColor(50, 150, 50)
        
        for i, item in enumerate(items, 1):
            p_item = doc.add_paragraph()
            p_item.paragraph_format.left_indent = Inches(0.6)
            p_item.paragraph_format.space_before = Pt(0)
            p_item.paragraph_format.space_after = Pt(1)
            p_item.paragraph_format.line_spacing = 1.1
            run_item = p_item.add_run(f"{i}. {item}")
            run_item.font.size = Pt(9.5)
            run_item.font.color.rgb = color_text_dark

    add_pdca_detail(
        "Identified Issues (2021 PDCA Review)", 
        [
            "High frequency of unplanned/urgent orders from customers (major driver of transport cost overrun).",
            "Inefficient double-truck usage (one for delivery, and another for picking up non-conforming/NG products).",
            "Supplier delivery delays and failure to follow the Pick-Up (PU) plan."
        ],
        is_issue=True
    )
    add_pdca_detail(
        "Implemented Solutions", 
        [
            "Organized alignment meetings with concerned departments to establish cooperative measures to decrease unplanned customer orders.",
            "Enforced closer tracking and proactive follow-ups with suppliers regarding the Pick-Up (PU) plan.",
            "Initiated on-site Quality Control (QC) visits to suppliers to resolve NG product issues at the source."
        ],
        is_issue=False
    )

    # --- Job 2 Cont'd (2020) ---
    add_job_header("Warehouse Foreman (WH Foreman) - Promoted from Leader", "2020", "Oizuru Chugen Packaging System Co., Ltd.")
    add_bullet("Supervised the Picking and Loading teams; supported the KIKAN system (ERP/Winsped) for outgoing inspection.")
    add_bullet("Assisted the supervisor/manager in achieving the Year 2020 KPI: keeping picking mistakes under 100 PPM.")
    
    add_pdca_detail(
        "Key Challenges & Root Causes (End of Year actual: 227 PPM)",
        [
            "Operators relying on memory for picking parts instead of referencing the physical tag labels.",
            "Lack of careful cross-checking by team leaders, especially for items in poly bags."
        ],
        is_issue=True
    )
    add_pdca_detail(
        "Implemented Solutions",
        [
            "Mandated picking teams based on specific zones/destinations (TTKL, Okamoto, Amata, Rayong).",
            "Required strict usage of picking lists to mark and check off items one by one.",
            "Instructed operators to strictly pick items by tag label rather than memory to avoid errors.",
            "Enforced pack-by-pack count audits for poly bag items by team leaders.",
            "Established a tracking system for picking mistakes per operator to drive continuous improvement."
        ],
        is_issue=False
    )

    # --- Job 3 ---
    add_job_header("Warehouse Leader", "2018 - 2019", "Oizuru Chugen Packaging System Co., Ltd.")
    add_bullet("Managed warehouse picking and loading operations; prepared raw materials for production lines.")
    add_bullet("Prepared, reviewed, and updated Work Instructions (WI) for presentation to the Warehouse Manager.")
    add_bullet("Assisted the Lead Auditor in conducting annual ISO 9001 and ISO 14001 internal audits, specifically auditing the Logistics department.")
    add_bullet("Led Kaizen and warehouse location improvement projects (re-organizing layouts and defining sub-locations).")
    add_bullet("Managed daily plans, prepared picking lists, and verified operator completions.")
    add_bullet("Conducted inventory management and regular stock-on-hand audits.")
    add_bullet("Completed training certifications: Leadership (Extraordinary Leader by TICA, Sep 2017), Forklift Operation (Oct 2017), and ISO 9001 / ISO 14001 Internal Auditor (Feb 2017).")

    # --- Job 4 ---
    add_job_header("Warehouse Leader (Receiving Control)", "Feb 2016 - 2017", "Oizuru Chugen Packaging System Co., Ltd.")
    add_bullet("Managed the incoming goods receiving control process in coordination with the purchasing section.")
    add_bullet("Monitored receiving processes, verifying incoming goods quantities against purchase documentation.")

    # 3. Education
    add_section_header("Education")
    p_edu = doc.add_paragraph()
    p_edu.paragraph_format.space_before = Pt(4)
    p_edu.paragraph_format.space_after = Pt(10)
    run_edu_title = p_edu.add_run("Bachelor of Public Administration\n")
    run_edu_title.font.bold = True
    run_edu_title.font.size = Pt(11)
    run_edu_title.font.color.rgb = color_primary
    
    run_edu_detail = p_edu.add_run("Sukhothai Thammathirat Open University  |  Graduated: May 2012")
    run_edu_detail.font.size = Pt(10)
    run_edu_detail.font.color.rgb = color_text_dark

    # 4. Professional Training & Certifications
    add_section_header("Professional Training & Certifications")
    add_bullet("Forklift Operation Certification (Annual refreshers completed)")
    add_bullet("ISO 9001 & ISO 14001 Quality & Environmental Management Systems (Internal Auditor)")
    add_bullet("Leadership Development: Extraordinary Leader (Certified by TICA)")

    # 5. Skills & Additional Information
    add_section_header("Skills & Additional Information")
    
    table_skills = doc.add_table(rows=3, cols=2)
    table_skills.autofit = True
    table_skills.allow_autofit = True
    
    skills_data = [
        ("Languages:", "Thai (Native)  |  English (Conversational / Fair)"),
        ("Computer Skills:", "Microsoft Excel (SUMIF, VLOOKUP, formulas), Microsoft PowerPoint, Microsoft Word"),
        ("ERP Systems:", "KIKAN System (similar to ERP, Winsped)")
    ]
    
    for i, (label, val) in enumerate(skills_data):
        row = table_skills.rows[i]
        cell_lbl = row.cells[0]
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.paragraph_format.space_after = Pt(4)
        run_lbl = p_lbl.add_run(label)
        run_lbl.font.bold = True
        run_lbl.font.size = Pt(10)
        run_lbl.font.color.rgb = color_primary
        
        cell_val = row.cells[1]
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.space_after = Pt(4)
        run_val = p_val.add_run(val)
        run_val.font.size = Pt(10)
        run_val.font.color.rgb = color_text_dark

    output_path = "/Users/Jeff/Desktop/Nattawat_Thonmon_Resume.docx"
    doc.save(output_path)
    print(f"Word document updated successfully at: {output_path}")

if __name__ == "__main__":
    build_resume()
